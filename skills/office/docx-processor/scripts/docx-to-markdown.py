#!/usr/bin/env python3
"""
Word 文档转 Markdown 脚本
用法: python docx-to-markdown.py <docx文件路径> [输出文件路径]
"""

import sys
import os
import re

try:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    print("请安装依赖: pip install python-docx")
    sys.exit(1)


def get_heading_level(paragraph) -> int:
    """获取标题级别"""
    style_name = paragraph.style.name.lower()
    if style_name.startswith("heading"):
        try:
            return int(style_name.replace("heading ", "").replace("heading", ""))
        except ValueError:
            pass
    if style_name == "title":
        return 1
    return 0


def process_runs(paragraph) -> str:
    """处理段落中的格式"""
    text = ""
    for run in paragraph.runs:
        run_text = run.text
        if not run_text:
            continue

        # 处理格式
        if run.bold and run.italic:
            run_text = f"***{run_text}***"
        elif run.bold:
            run_text = f"**{run_text}**"
        elif run.italic:
            run_text = f"*{run_text}*"

        if run.underline:
            run_text = f"<u>{run_text}</u>"

        text += run_text

    return text


def is_list_item(paragraph) -> tuple:
    """检查是否是列表项"""
    style_name = paragraph.style.name.lower()

    # 检查样式名
    if "list" in style_name:
        if "number" in style_name or "ordered" in style_name:
            return (True, "ordered")
        return (True, "unordered")

    # 检查段落格式
    if hasattr(paragraph, "_p") and paragraph._p.pPr is not None:
        numPr = paragraph._p.pPr.numPr
        if numPr is not None:
            return (True, "ordered")

    return (False, None)


def process_table(table) -> str:
    """处理表格"""
    markdown = "\n"

    for i, row in enumerate(table.rows):
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        markdown += "| " + " | ".join(cells) + " |\n"

        # 添加表头分隔行
        if i == 0:
            markdown += "| " + " | ".join(["---"] * len(cells)) + " |\n"

    return markdown + "\n"


def docx_to_markdown(docx_path: str) -> str:
    """将 Word 文档转换为 Markdown"""
    doc = Document(docx_path)
    markdown = ""
    list_counter = 0
    in_list = False

    for element in doc.element.body:
        # 处理段落
        if element.tag.endswith("p"):
            for paragraph in [p for p in doc.paragraphs if p._p == element]:
                text = process_runs(paragraph)

                if not text.strip():
                    if in_list:
                        in_list = False
                        list_counter = 0
                    markdown += "\n"
                    continue

                # 处理标题
                heading_level = get_heading_level(paragraph)
                if heading_level > 0:
                    markdown += f"\n{'#' * heading_level} {text}\n\n"
                    continue

                # 处理列表
                is_list, list_type = is_list_item(paragraph)
                if is_list:
                    in_list = True
                    if list_type == "ordered":
                        list_counter += 1
                        markdown += f"{list_counter}. {text}\n"
                    else:
                        markdown += f"- {text}\n"
                    continue
                else:
                    if in_list:
                        in_list = False
                        list_counter = 0
                        markdown += "\n"

                # 处理对齐
                if paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                    text = f"<center>{text}</center>"
                elif paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
                    text = f'<div align="right">{text}</div>'

                markdown += f"{text}\n\n"

        # 处理表格
        elif element.tag.endswith("tbl"):
            for table in doc.tables:
                if table._tbl == element:
                    markdown += process_table(table)
                    break

    # 清理多余空行
    markdown = re.sub(r"\n{3,}", "\n\n", markdown)

    return markdown.strip()


def main():
    if len(sys.argv) < 2:
        print("用法: python docx-to-markdown.py <docx文件路径> [输出文件路径]")
        sys.exit(1)

    docx_path = sys.argv[1]

    if not os.path.exists(docx_path):
        print(f"文件不存在: {docx_path}")
        sys.exit(1)

    # 确定输出路径
    if len(sys.argv) > 2:
        output_path = sys.argv[2]
    else:
        output_path = os.path.splitext(docx_path)[0] + ".md"

    print(f"📄 转换 Word 文档: {docx_path}")

    markdown = docx_to_markdown(docx_path)

    with open(output_path, "w", encoding="utf-8") as f:
        f.write(markdown)

    print(f"✅ Markdown 已保存: {output_path}")
    print(f"   字数: {len(markdown)} 字符")


if __name__ == "__main__":
    main()
